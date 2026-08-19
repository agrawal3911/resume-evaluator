import io
from docx import Document
from app.parsers.cleaning import clean_text
from app.parsers.exceptions import DocumentParsingError

def parse_docx(file_bytes: bytes) -> str:
    """
    Parses a DOCX from bytes and extracts its text.
    Raises DocumentParsingError if the document is invalid or empty.
    """
    try:
        file_stream = io.BytesIO(file_bytes)
        doc = Document(file_stream)
    except Exception as e:
        raise DocumentParsingError(f"Failed to open DOCX document: {e}") from e

    extracted_text: list[str] = []

    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text:
            extracted_text.append(para.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    extracted_text.append(cell.text.replace("\n", " "))

    # Combine, clean, then validate
    combined_text = "\n".join(extracted_text)
    cleaned = clean_text(combined_text)

    if not cleaned:
        raise DocumentParsingError(
            "No extractable text found in DOCX. The document may be empty."
        )

    return cleaned
