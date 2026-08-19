import pytest
import io
import pymupdf as fitz
from docx import Document
from app.parsers.cleaning import clean_text
from app.parsers.exceptions import DocumentParsingError, UnsupportedFormatError
from app.parsers.pdf_parser import parse_pdf
from app.parsers.docx_parser import parse_docx
from app.parsers.document_parser import parse_document

# --- Helper functions to create in-memory documents ---

def create_valid_pdf_bytes(text: str) -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), text)
    return doc.write()

def create_empty_pdf_bytes() -> bytes:
    doc = fitz.open()
    doc.new_page()
    return doc.write()

def create_valid_docx_bytes(text: str) -> bytes:
    doc = Document()
    doc.add_paragraph(text)
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()

def create_empty_docx_bytes() -> bytes:
    doc = Document()
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()

# --- Tests for cleaning.py ---

def test_clean_text():
    raw_text = "  This   is \n\n\n\n a test   \n  "
    expected = "This is\n\na test"
    assert clean_text(raw_text) == expected
    assert clean_text("") == ""
    assert clean_text(None) == ""

# --- Tests for pdf_parser.py ---

def test_parse_valid_pdf():
    pdf_bytes = create_valid_pdf_bytes("Hello PDF!")
    extracted = parse_pdf(pdf_bytes)
    assert "Hello PDF!" in extracted

def test_parse_empty_pdf():
    pdf_bytes = create_empty_pdf_bytes()
    with pytest.raises(DocumentParsingError, match="No extractable text"):
        parse_pdf(pdf_bytes)

def test_parse_invalid_pdf():
    with pytest.raises(DocumentParsingError, match="Failed to open PDF"):
        parse_pdf(b"Not a real PDF")

# --- Tests for docx_parser.py ---

def test_parse_valid_docx():
    docx_bytes = create_valid_docx_bytes("Hello DOCX!")
    extracted = parse_docx(docx_bytes)
    assert "Hello DOCX!" in extracted

def test_parse_empty_docx():
    docx_bytes = create_empty_docx_bytes()
    with pytest.raises(DocumentParsingError, match="No extractable text"):
        parse_docx(docx_bytes)

def test_parse_invalid_docx():
    with pytest.raises(DocumentParsingError, match="Failed to open DOCX"):
        parse_docx(b"Not a real DOCX")

# --- Tests for document_parser.py ---

def test_parse_document_routing():
    pdf_bytes = create_valid_pdf_bytes("PDF Content")
    docx_bytes = create_valid_docx_bytes("DOCX Content")
    
    assert "PDF Content" in parse_document(pdf_bytes, "resume.pdf")
    assert "PDF Content" in parse_document(pdf_bytes, "resume.PDF") # case insensitive
    assert "DOCX Content" in parse_document(docx_bytes, "resume.docx")

def test_parse_document_unsupported():
    with pytest.raises(UnsupportedFormatError, match="Unsupported file format"):
        parse_document(b"random bytes", "resume.txt")

# --- Additional tests ---

def test_parse_multi_page_pdf():
    doc = fitz.open()
    page1 = doc.new_page()
    page1.insert_text((50, 50), "Page one content")
    page2 = doc.new_page()
    page2.insert_text((50, 50), "Page two content")
    pdf_bytes = doc.write()
    doc.close()

    extracted = parse_pdf(pdf_bytes)
    assert "Page one content" in extracted
    assert "Page two content" in extracted

def test_parse_whitespace_only_pdf():
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "   \n\n   \n  ")
    pdf_bytes = doc.write()
    doc.close()

    with pytest.raises(DocumentParsingError, match="No extractable text"):
        parse_pdf(pdf_bytes)

def test_parse_multi_paragraph_docx_with_table():
    doc = Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Cell A"
    table.cell(0, 1).text = "Cell B"
    stream = io.BytesIO()
    doc.save(stream)
    docx_bytes = stream.getvalue()

    extracted = parse_docx(docx_bytes)
    assert "First paragraph" in extracted
    assert "Second paragraph" in extracted
    assert "Cell A" in extracted
    assert "Cell B" in extracted

def test_parse_whitespace_only_docx():
    doc = Document()
    doc.add_paragraph("   \n  \n   ")
    stream = io.BytesIO()
    doc.save(stream)
    docx_bytes = stream.getvalue()

    with pytest.raises(DocumentParsingError, match="No extractable text"):
        parse_docx(docx_bytes)

def test_parse_document_uppercase_extensions():
    pdf_bytes = create_valid_pdf_bytes("Upper PDF")
    docx_bytes = create_valid_docx_bytes("Upper DOCX")

    assert "Upper PDF" in parse_document(pdf_bytes, "resume.PDF")
    assert "Upper DOCX" in parse_document(docx_bytes, "resume.DOCX")

def test_parse_document_unsupported_various():
    with pytest.raises(UnsupportedFormatError):
        parse_document(b"fake", "photo.jpg")
    with pytest.raises(UnsupportedFormatError):
        parse_document(b"fake", "page.html")
